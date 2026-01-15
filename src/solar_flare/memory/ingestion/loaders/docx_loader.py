"""
Word document loader using python-docx.
"""
from pathlib import Path
from typing import BinaryIO, Dict, Any

from .base import BaseDocumentLoader, LoadedDocument, DocumentLoadError


class DocxLoader(BaseDocumentLoader):
    """Load Word documents using python-docx."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def __init__(self, include_headers_footers: bool = True, include_tables: bool = True):
        """
        Initialize Word loader.

        Args:
            include_headers_footers: Whether to extract headers/footers
            include_tables: Whether to extract table content
        """
        self.include_headers_footers = include_headers_footers
        self.include_tables = include_tables

    def load(self, file_path: Path) -> LoadedDocument:
        """Load Word document from file path."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentLoadError(
                "python-docx not installed. Run: pip install python-docx",
                str(file_path)
            )

        try:
            doc = Document(file_path)
            return self._extract_content(doc, str(file_path), file_path.name)
        except Exception as e:
            raise DocumentLoadError(f"Failed to load Word document: {e}", str(file_path), e)

    def load_from_stream(
        self,
        stream: BinaryIO,
        filename: str,
        source_path: str,
    ) -> LoadedDocument:
        """Load Word document from binary stream."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentLoadError(
                "python-docx not installed. Run: pip install python-docx",
                source_path
            )

        try:
            doc = Document(stream)
            return self._extract_content(doc, source_path, filename)
        except Exception as e:
            raise DocumentLoadError(f"Failed to load Word stream: {e}", source_path, e)

    def _extract_content(self, doc, source_path: str, filename: str) -> LoadedDocument:
        """Extract content from python-docx Document."""
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Include heading style if present
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.replace("Heading ", "").strip())
                        text_parts.append(f"{'#' * level} {para.text}")
                    except ValueError:
                        text_parts.append(para.text)
                else:
                    text_parts.append(para.text)

        # Extract tables
        if self.include_tables:
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")

        content = "\n\n".join(text_parts)

        # Extract core properties if available
        metadata: Dict[str, Any] = {"filename": filename}
        try:
            if doc.core_properties:
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.created:
                    metadata["created"] = str(doc.core_properties.created)
                if doc.core_properties.modified:
                    metadata["modified"] = str(doc.core_properties.modified)
                if doc.core_properties.title:
                    metadata["doc_title"] = doc.core_properties.title
        except Exception:
            pass

        title = metadata.get("doc_title") or self._generate_title(filename)

        return LoadedDocument(
            content=content,
            title=title,
            source_path=source_path,
            file_format="docx",
            content_hash=LoadedDocument.compute_hash(content),
            metadata=metadata,
        )
